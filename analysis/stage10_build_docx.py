# -*- coding: utf-8 -*-
"""Rebuild the manuscript docx from scratch from the current canonical sources:
main.tex (title/abstract/table/figure/backmatter) + the 6 .md body files
(Introduction/Methods/Results/Discussion/Limitations/Conclusion prose).

Why this script exists (Stage 10, 2026-07-10; rebuilt 2026-07-11 for the Plan-A-v2
"reviewer-accepted restructure" per user-supplied restructuring docx): the manuscript.docx
deliverable found in this folder at the start of Stage 10 was not merely stale by a few
hours -- it was the ENTIRE PRE-RESTRUCTURING manuscript. A docx built by hand or patched
incrementally is exactly how that drift happens silently. This script is the SOLE canonical
path to the docx: it reads only main.tex + the 6 .md files (never edits a previous docx), so
re-running it after any manuscript edit regenerates a docx guaranteed consistent with the
current LaTeX/markdown source of truth. Re-run this script (not a manual Word edit) any
time the manuscript content changes.

2026-07-11 restructure notes: dropped Research-in-Context panel (not in the new structure),
dropped the structural-break Table/worked-forecast Table entirely (demoted to a disclosed
Limitations item per the accepted restructuring plan), added a new eligibility-tier Table 2,
renumbered Tables 2-3 -> 3-4, added Figures 3-4 (order-selection-effect chart; walk-forward
chart), added standalone Limitations and Conclusion sections read from their own .md files
(previously Conclusion was inline main.tex text -- now built via conclusion_body.tex), added
5 new citations (Akaike 1974, Hurvich and Tsai 1989, Perla 2021, Siami-Namini 2018, Jain 2024,
Shoko 2026, Guleryuz 2021, Osman 2026), then a further round (2026-07-11, same session, per
user follow-up) restoring Ayisah/Piu/Lasim and adding 10 more real, verified citations
(Box/Jenkins, Hochreiter and Schmidhuber, Kingma and Ba, Seabold and Perktold, Paszke et al,
Burnham and Anderson, von Elm et al/STROBE, Hyndman and Athanasopoulos, UN IGME, Foreman et
al/GBD forecasting) to reach 25 total, each integrated at a genuine, load-bearing point in the
prose rather than appended as an orphan.

Known residual limitation: no pdflatex/tectonic/pandoc is available in this environment,
so the LaTeX itself has never been compile-tested -- only structurally audited (brace
balance, \\cite/\\ref/\\label cross-checks, table column counts). Recommended before
submission: upload the latex/ folder to Overleaf (free tier) or run a local TeX
distribution to confirm it compiles cleanly.
"""
import re
import csv
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = r"C:\Users\VGhanem\Documents\Claude\Projects\Public Health & Epidemiology Research Skills\23. Disease Burden Forecasting Ghana"
MS = ROOT + r"\manuscript"
LX = MS + r"\latex"

def read(path):
    with open(path, encoding="utf-8") as f:
        return f.read()

main_tex = read(LX + r"\main.tex")

TITLE = ("A reproducible workflow and source-provenance audit for forecasting short "
         "national health-indicator panels: a worked example from Ghana using WHO Global "
         "Health Observatory and World Bank data")

# ---------- helpers ----------
def strip_tex(t):
    """Pre-process a raw main.tex substring: convert \\cite{key} -> [key] (so the same
    bracket-to-number tokenizer used everywhere else picks it up), \\ref{label} -> REF:label,
    convert inline math-mode symbols actually used in this manuscript's Abstract to their
    plain-text equivalents, and unescape %, _, --, ~. Must be called on every main.tex-sourced
    block before add_para.

    BUG FIX (Phase-12 QA gate, manuscript editor review, 2026-07-11): this function previously
    left math-mode LaTeX ($\\Delta$AICc $<$ 2, $\\sim$24\\%) completely unconverted -- it leaked
    into the docx as raw source text -- and blindly replaced every literal '~' with a space,
    which is correct for LaTeX's "Table~\\ref{}" non-breaking-space idiom but wrong for '~' used
    informally as an approx-equals sign, producing a silently-vanished tilde and a stray double
    space ("up to  24%"). The '~' bug is now fixed at the source instead (main.tex uses the
    proper $\\sim$ math command), and this function converts that math command explicitly rather
    than deleting it.
    """
    t = re.sub(r"\\cite[a-zA-Z]*\{([^}]+)\}", lambda m: "[" + m.group(1) + "]", t)
    t = re.sub(r"\\ref\{([^}]+)\}", lambda m: "REF:" + m.group(1), t)
    t = t.replace(r"$\Delta$", "Δ").replace(r"$\sim$", "≈")
    t = t.replace(r"$<$", "<").replace(r"$>$", ">")
    t = t.replace(r"$\geq$", "≥").replace(r"$\leq$", "≤")
    t = t.replace(r"\%", "%").replace(r"\_", "_").replace(r"\&", "&").replace(r"\#", "#")
    t = t.replace("--", "–").replace("~", " ")
    return t

def add_runs_from_markup(paragraph, text, cite_map):
    """Parse **bold**, \\textbf{}, \\textit{}, [key] citation markers, REF: markers into runs."""
    text = re.sub(r"\\textbf\{([^}]*)\}", r"**\1**", text)
    text = re.sub(r"\\textit\{([^}]*)\}", r"_\1_", text)
    text = re.sub(r"\\texttt\{([^}]*)\}", r"\1", text)
    token_re = re.compile(r"(\*\*.+?\*\*|_[^_]+_|\[[a-zA-Z0-9_,\s]+\]|REF:[a-zA-Z0-9:_]+)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 1:
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("[") and part.endswith("]") and re.match(r"^\[[a-zA-Z0-9_,\s]+\]$", part):
            keys = [k.strip() for k in part[1:-1].split(",")]
            nums = [str(cite_map[k]) for k in keys if k in cite_map]
            if nums:
                r = paragraph.add_run("[" + ",".join(nums) + "]")
                r.font.superscript = False
            else:
                paragraph.add_run(part)
        elif part.startswith("REF:tab:table"):
            # BUG FIX (Phase-12 QA gate, manuscript editor review, 2026-07-11): every \ref{} in
            # this manuscript is already preceded by the literal word "Table"/"Figure" in the
            # surrounding prose (verified across all .tex source files -- there is no bare \ref{}
            # usage anywhere), so prepending it again here produced "Table Table 1" in the
            # rendered Abstract. Emit the bare number only.
            paragraph.add_run(part[len("REF:tab:table"):])
        elif part.startswith("REF:fig:figure"):
            paragraph.add_run(part[len("REF:fig:figure"):])
        else:
            paragraph.add_run(part)

def add_para(doc, text, cite_map, style=None):
    p = doc.add_paragraph(style=style)
    add_runs_from_markup(p, text, cite_map)
    return p

# ---------- bib entries (parsed manually against references.bib content) ----------
BIB = {
"aryee2018": "Aryee G, Kwarteng E, Essuman R, Nkansa Agyei A, Kudzawu S, Djagbletey R, Owusu Darkwa E, Forson A. Estimating the incidence of tuberculosis cases reported at a tertiary hospital in Ghana: a time series model approach. BMC Public Health. 2018;18(1):1292. doi:10.1186/s12889-018-6221-z.",
"inusah2026": "Inusah AW, Jagha T, Head MG, Seidu AA, Ziblim SD. Trends, regional disparities, and projected burden of anemia among women of reproductive age in Ghana, 2000-2030. Reprod Health. 2026;23(1). doi:10.1186/s12978-026-02345-6.",
"hyndman2008": "Hyndman RJ, Khandakar Y. Automatic time series forecasting: the forecast package for R. J Stat Softw. 2008;27(3):1-22. doi:10.18637/jss.v027.i03.",
"akaike1974": "Akaike H. A new look at the statistical model identification. IEEE Trans Autom Control. 1974;19(6):716-723. doi:10.1109/TAC.1974.1100705.",
"hurvich1989": "Hurvich CM, Tsai CL. Regression and time series model selection in small samples. Biometrika. 1989;76(2):297-307. doi:10.1093/biomet/76.2.297.",
"perla2021": "Perla F, Richman R, Scognamiglio S, Wuthrich MV. Time-series forecasting of mortality rates using deep learning. Scand Actuar J. 2021. doi:10.1080/03461238.2020.1867232.",
"nikparvar2021": "Nikparvar B, Rahman MM, Hatami F, Thill JC. Spatio-temporal prediction of the COVID-19 pandemic in US counties: modeling with a deep LSTM neural network. Sci Rep. 2021;11:21715.",
"siaminamini2018": "Siami-Namini S, Tavakoli N, Namin AS. A comparison of ARIMA and LSTM in forecasting time series. In: 2018 17th IEEE ICMLA. IEEE; 2018:1394-1401. doi:10.1109/ICMLA.2018.00227.",
"jain2024": "Jain S, Agrawal S, Mohapatra E, et al. A novel ensemble ARIMA-LSTM approach for evaluating COVID-19 cases and future outbreak preparedness. Health Care Sci. 2024;3(6):409-425. doi:10.1002/hcs2.123.",
"shoko2026": "Shoko C, Sigauke C, Makatjane K. Hierarchical forecasting of COVID-19 cases in Africa using machine learning models. Front Epidemiol. 2026;6:1696282. doi:10.3389/fepid.2026.1696282.",
"guleryuz2021": "Guleryuz D. Forecasting outbreak of COVID-19 in Turkey; comparison of Box-Jenkins, Brown's exponential smoothing and long short-term memory models. Process Saf Environ Prot. 2021;149:927-935. doi:10.1016/j.psep.2021.03.032.",
"osman2026": "Osman AA, Mengistie DT, Adawe DH, Figa RT, Marine BT. A comparative approach to analyzing and forecasting carbon dioxide emissions in Ethiopia using Bayesian ARIMA and Bayesian structural time series models. Carbon Balance Manag. 2026. doi:10.1186/s13021-026-00480-y.",
"box2015": "Box GEP, Jenkins GM, Reinsel GC, Ljung GM. Time Series Analysis: Forecasting and Control. 5th ed. Hoboken, NJ: Wiley; 2015.",
"hochreiter1997": "Hochreiter S, Schmidhuber J. Long short-term memory. Neural Comput. 1997;9(8):1735-1780. doi:10.1162/neco.1997.9.8.1735.",
"kingma2015": "Kingma DP, Ba J. Adam: a method for stochastic optimization. 3rd International Conference on Learning Representations (ICLR); 2015.",
"seabold2010": "Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. Proceedings of the 9th Python in Science Conference (SciPy); 2010:92-96.",
"paszke2019": "Paszke A, Gross S, Massa F, et al. PyTorch: an imperative style, high-performance deep learning library. Advances in Neural Information Processing Systems 32 (NeurIPS); 2019:8024-8035.",
"burnham2002": "Burnham KP, Anderson DR. Model Selection and Multimodel Inference: A Practical Information-Theoretic Approach. 2nd ed. New York, NY: Springer; 2002.",
"vonelm2007": "von Elm E, Altman DG, Egger M, Pocock SJ, Gotzsche PC, Vandenbroucke JP. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet. 2007;370(9596):1453-1457. doi:10.1016/S0140-6736(07)61602-X.",
"hyndman2021fpp": "Hyndman RJ, Athanasopoulos G. Forecasting: Principles and Practice. 3rd ed. Melbourne: OTexts; 2021.",
"unigme2023": "UN Inter-agency Group for Child Mortality Estimation. Levels and Trends in Child Mortality: Report 2023. UNICEF, WHO, World Bank Group, UN DESA Population Division; 2023.",
"foreman2024": "Foreman KJ, et al. Burden of disease scenarios for 204 countries and territories, 2022-2050: a forecasting analysis for the Global Burden of Disease Study 2021. Lancet. 2024. doi:10.1016/S0140-6736(24)00685-8.",
"ayisah2025": "Ayisah C, et al. Quality of routine malaria data captured at primary health facilities in the Hohoe Municipality, Ghana. Sci Rep. 2025.",
"piu2024": "Piu LJ, et al. Assessment of expanded programme on immunization routine data quality in the upper east region of Ghana. BMC Health Serv Res. 2024;24(1):886. doi:10.1186/s12913-024-11347-8.",
"lasim2022": "Lasim OU, et al. Maternal and child health data quality in health care facilities at the Cape Coast Metropolis, Ghana. BMC Health Serv Res. 2022.",
}

# The .md files (canonical prose source for Intro/Methods/Results/Discussion) cite in plain
# "[Author et al., Year]" form, not \cite{key}. Multi-citation brackets are semicolon-separated
# (e.g. "[Perla et al., 2021; Nikparvar et al., 2021]") and are split before lookup.
MD_CITE_TEXT_TO_KEY = {
    "Aryee et al., 2018": "aryee2018",
    "Inusah et al., 2026": "inusah2026",
    "Hyndman & Khandakar, 2008": "hyndman2008",
    "Akaike, 1974": "akaike1974",
    "Hurvich & Tsai, 1989": "hurvich1989",
    "Perla et al., 2021": "perla2021",
    "Nikparvar et al., 2021": "nikparvar2021",
    "Siami-Namini et al., 2018": "siaminamini2018",
    "Jain et al., 2024": "jain2024",
    "Shoko et al., 2026": "shoko2026",
    "Guleryuz, 2021": "guleryuz2021",
    "Osman et al., 2026": "osman2026",
    "Box et al., 2015": "box2015",
    "Hochreiter & Schmidhuber, 1997": "hochreiter1997",
    "Kingma & Ba, 2015": "kingma2015",
    "Seabold & Perktold, 2010": "seabold2010",
    "Paszke et al., 2019": "paszke2019",
    "Burnham & Anderson, 2002": "burnham2002",
    "von Elm et al., 2007": "vonelm2007",
    "Hyndman & Athanasopoulos, 2021": "hyndman2021fpp",
    "UN IGME, 2023": "unigme2023",
    "Foreman et al., 2024": "foreman2024",
    "Piu et al., 2024": "piu2024",
    "Lasim et al., 2022": "lasim2022",
    "Ayisah et al., 2025": "ayisah2025",
}

# ---------- citation order (true first-appearance order across reading order) ----------
# Reading order in the rendered docx: Introduction -> Methods -> Results -> Discussion ->
# Limitations (which restores the DHIMS2 data-quality citations cut from Discussion for length).
# Order is extracted from the .md files themselves, since that is what add_md_section()
# actually renders into the docx.
def extract_cite_order():
    seen = []
    bracket_re = re.compile(r"\[([^\[\]]+)\]")
    for fn in ["introduction.md", "methods.md", "results.md", "discussion.md", "limitations.md"]:
        txt = read(MS + "\\" + fn)
        for m in bracket_re.finditer(txt):
            inner = m.group(1)
            if not re.search(r"\d{4}", inner):
                continue
            for part in inner.split(";"):
                key = MD_CITE_TEXT_TO_KEY.get(part.strip())
                if key and key not in seen:
                    seen.append(key)
    return seen

cite_order = extract_cite_order()
cite_map = {k: i + 1 for i, k in enumerate(cite_order)}

assert set(BIB) == set(cite_order), "BIB dict keys must exactly match cite_order keys: " + str(set(BIB) ^ set(cite_order))
print("Citation count check: %d cited, %d bib entries -- match: %s" % (len(cite_order), len(BIB), set(BIB) == set(cite_order)))

def md_citations_to_numbered(text):
    def repl(m):
        inner = m.group(1)
        if not re.search(r"\d{4}", inner):
            return m.group(0)
        nums = []
        for part in inner.split(";"):
            key = MD_CITE_TEXT_TO_KEY.get(part.strip())
            if key is None:
                return m.group(0)  # leave untouched if not a recognised citation bracket
            nums.append(str(cite_map[key]))
        return "[" + ",".join(nums) + "]"
    return re.sub(r"\[([^\[\]]+)\]", repl, text)

# ---------- build document ----------
doc = docx.Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = docx.oxml.OxmlElement('w:rFonts')
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Times New Roman')

for i in range(1, 4):
    hstyle = doc.styles[f"Heading {i}"]
    hstyle.font.name = "Times New Roman"
    hstyle.font.color.rgb = RGBColor(0, 0, 0)
    hstyle.font.bold = True
    hstyle.font.size = Pt(14) if i == 1 else Pt(12.5) if i == 2 else Pt(12)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(TITLE)
r.bold = True
r.font.size = Pt(15)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Valentine Golden Ghanem")
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run("Ghana COCOBOD Cocoa Clinic, Accra, Ghana")
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run("ORCID: 0009-0002-8332-0220 | Correspondence: valentineghanem@gmail.com")
doc.add_paragraph()

# Abstract
doc.add_heading("Abstract", level=1)
abs_block = main_tex[main_tex.find(r"\begin{abstract}") + len(r"\begin{abstract}"):main_tex.find(r"\end{abstract}")]
for para_txt in re.split(r"\n\n+", abs_block.strip()):
    para_txt = para_txt.strip()
    if not para_txt:
        continue
    add_para(doc, strip_tex(para_txt), cite_map)

kw_block = main_tex[main_tex.find(r"\begin{keyword}") + len(r"\begin{keyword}"):main_tex.find(r"\end{keyword}")]
kws = [k.strip() for k in kw_block.replace("\n", " ").split(r"\sep")]
p = doc.add_paragraph()
r = p.add_run("Keywords: ")
r.bold = True
p.add_run("; ".join(kws))

# Introduction / Methods / Results (source: .md files, canonical prose)
def add_md_section(doc, md_path, cite_map, skip_first_h1=True):
    txt = read(md_path)
    blocks = txt.split("\n\n")
    first = True
    for block in blocks:
        block = block.strip("\n")
        if not block.strip():
            continue
        if block.startswith("# "):
            if skip_first_h1 and first:
                first = False
                continue
            doc.add_heading(block[2:].strip(), level=1)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith("### "):
            doc.add_heading(block[4:].strip(), level=3)
        elif block.startswith("**Figure"):
            cap = doc.add_paragraph()
            m = re.match(r"\*\*(Figure \d+\.)\*\*\s*(.*)", block, re.DOTALL)
            if m:
                r = cap.add_run(m.group(1) + " ")
                r.bold = True
                cap.add_run(m.group(2))
            else:
                add_para(doc, block, cite_map)
        else:
            block = md_citations_to_numbered(block)
            add_para(doc, block, cite_map)
        first = False

doc.add_heading("Introduction", level=1)
add_md_section(doc, MS + r"\introduction.md", cite_map, skip_first_h1=True)

doc.add_heading("Methods", level=1)
add_md_section(doc, MS + r"\methods.md", cite_map, skip_first_h1=True)

doc.add_heading("Results", level=1)
add_md_section(doc, MS + r"\results.md", cite_map, skip_first_h1=True)

# Table 1 (documented data-integrity corrections)
doc.add_paragraph()
cap = doc.add_paragraph()
r = cap.add_run("Table 1. ")
r.bold = True
cap.add_run("Documented data-integrity corrections required to build an analysis-ready national indicator panel from the WHO GHO/GHE and World Bank source files used in this paper.")

t1_headers = ["Source file / series", "Pitfall", "Symptom if uncorrected", "Correction applied", "Verification method"]
t1_rows = [
    ["All WHO GHO/GHE exports", "Leading hashtag (HXL) metadata row",
     "Non-numeric row corrupts type coercion; silent NaNs or fit failure",
     "Drop HXL row before numeric coercion",
     "Structural -- fixed row position in WHO export format"],
    ["Under-five mortality rate", "Wealth-quintile stratification (Q1-Q5 + national total) presenting as year-level duplication",
     "Averaging across quintiles yields a value that is not the national rate",
     "Retained the national total series (WEALTHQUINTILE_TOTL); did not average across quintiles",
     "Confirmed directly against the WHO GHO OData API (ghoapi.azureedge.net)"],
    ["Air-pollution attributable mortality/DALY indicators (2 series)",
     "Cause-of-death sub-category stratification presenting as year-level duplication",
     "Summing sub-categories double-counts; averaging understates",
     "Retained the per-year maximum (the “all-causes” total)",
     "Confirmed by exact/near-exact equality of the maximum to the sum of its sub-category components"],
    ["Crude suicide rate, year 2021",
     "Recurring WHO GHO export anomaly (36 rows vs. expected 3; only 3 carry valid confidence intervals)",
     "Spurious rows distort the 2021 value and interval",
     "Retained only the confidence-interval-bearing rows",
     "Row-count vs. expected structure; CI presence as validity filter"],
    ["World Health Statistics compendium + one consolidated GHO export",
     "Near-total re-exports of the disease-specific source files",
     "Merging as independent observations double-counts",
     "Used exclusively as internal consistency cross-checks; never merged as additional observations",
     "Value-level comparison against four spot-checked indicators (exact match, 0.0 difference, on tuberculosis incidence, life expectancy, total NCD deaths, and NCD premature-mortality %)"],
    ["Current health expenditure (% GDP): WHO-GHO vs. World Bank",
     "Two “sources” numerically identical to several decimal places",
     "Reported as an independent cross-source check when it is not",
     "Flagged as non-independent (both trace to the WHO Global Health Expenditure Database); excluded from cross-source validation claims",
     "Decimal-level equality across the panel"],
]
t1 = doc.add_table(rows=1 + len(t1_rows), cols=len(t1_headers))
t1.style = "Light Grid Accent 1"
for j, h in enumerate(t1_headers):
    r = t1.cell(0, j).paragraphs[0].add_run(h)
    r.bold = True
for i, rowvals in enumerate(t1_rows, start=1):
    for j, v in enumerate(rowvals):
        t1.cell(i, j).text = v
for row in t1.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for r in p.runs:
                r.font.size = Pt(8.5)
doc.add_paragraph()

LABELS = {
"life_expectancy_birth_yrs_who": "Life expectancy at birth (WHO-GHO), yrs",
"hale_birth_yrs": "Healthy life expectancy at birth, yrs",
"u5mr_who": "Under-five mortality rate, per 1,000",
"ncd_premature_pct": "NCD premature mortality, %",
"ncd_3070_probability_pct": "NCD 30–70 death probability, %",
"total_ncd_deaths": "Total NCD deaths, n",
"suicide_rate_crude_who": "Suicide rate (crude, WHO-GHO), per 100,000",
"suicide_rate_agestd_who": "Suicide rate (age-std., WHO-GHO), per 100,000",
"tb_incidence_per100k": "Tuberculosis incidence, per 100,000",
"hiv_new_infections_n": "HIV new infections, n",
"hiv_new_infections_per1000": "HIV new infections, per 1,000",
"malaria_incidence_per1000atrisk": "Malaria incidence, per 1,000 at risk",
"oop_pct_che": "Out-of-pocket exp., % of CHE",
"oop_percapita_usd": "Out-of-pocket exp. per capita, USD",
"che_pct_gdp": "Current health exp., % GDP (WHO-GHO)",
"che_percapita_usd": "Current health exp. per capita, USD",
"life_expectancy_birth_yrs_wb": "Life expectancy at birth (World Bank), yrs",
"suicide_rate_wb": "Suicide rate (World Bank), per 100,000",
"che_pct_gdp_wb": "Current health exp., % GDP (World Bank)",
"tfr_national_wb": "Total fertility rate (World Bank), births/woman",
"pm25_exposure_wb": "PM2.5 exposure (World Bank), µg/m³",
}
ROW_ORDER = list(LABELS.keys())
NON_CONVERGENT = {"u5mr_who", "hiv_new_infections_n", "hiv_new_infections_per1000", "pm25_exposure_wb"}

rows = []
with open(ROOT + r"\outputs\data\national_forecasts_2030.csv", encoding="utf-8") as f:
    for row in csv.DictReader(f):
        rows.append(row)
rows_by_ind = {r["indicator"]: r for r in rows}

def fmt(v, dec=2):
    return f"{float(v):,.{dec}f}"

def fmt_int(v):
    return f"{float(v):,.0f}"

# Table 2 (NEW -- eligibility-tier list, 21 forecastable series)
cap = doc.add_paragraph()
r = cap.add_run("Table 2. ")
r.bold = True
cap.add_run("The 21 forecastable series by data-length confidence tier. Standard: n≥25; Low: 15-24 years. "
             "Four further indicators (UHC Service Coverage Index, n=7; three air-pollution indicators, n=10 each) "
             "fell below the 15-year formal-forecasting minimum and are retained as descriptive context only (not shown).")

t2_headers = ["Indicator", "n", "Last year", "Tier"]
t2 = doc.add_table(rows=1 + len(ROW_ORDER), cols=len(t2_headers))
t2.style = "Light Grid Accent 1"
for j, h in enumerate(t2_headers):
    r = t2.cell(0, j).paragraphs[0].add_run(h)
    r.bold = True
for i, ind in enumerate(ROW_ORDER, start=1):
    d = rows_by_ind[ind]
    tier = "Standard" if d["confidence"].strip().lower() == "standard" else "Low"
    vals = [LABELS[ind], d["n_obs"], d["last_year"], tier]
    for j, v in enumerate(vals):
        t2.cell(i, j).text = str(v)
for row in t2.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for r in p.runs:
                r.font.size = Pt(9)
doc.add_paragraph()

# Table 3 (walk-forward validation)
cap = doc.add_paragraph()
r = cap.add_run("Table 3. ")
r.bold = True
cap.add_run("Walk-forward validation of classical versus LSTM forecasting methods on two representative series (Methods). "
             "MAPE = mean absolute percentage error, seed-averaged across five independent LSTM initializations per fold "
             "and hidden-layer size; SD = across-seed standard deviation, rescaled to the same percentage-error units as "
             "MAPE (Figure 4 plots the identical rescaled quantity).")

wf_headers = ["Series", "n", "Folds", "ETS MAPE (%)", "ARIMA MAPE (%)", "LSTM h=8 MAPE (%)", "LSTM h=32 MAPE (%)"]
wf_rows = []
with open(ROOT + r"\outputs\data\walkforward_validation_results.csv", encoding="utf-8") as f:
    for row in csv.DictReader(f):
        wf_rows.append(row)
WF_LABELS = {"u5mr_who": "Under-five mortality (long series)", "tb_incidence_per100k": "Tuberculosis incidence (short series)"}
t3 = doc.add_table(rows=1 + len(wf_rows), cols=len(wf_headers))
t3.style = "Light Grid Accent 1"
for j, h in enumerate(wf_headers):
    r = t3.cell(0, j).paragraphs[0].add_run(h)
    r.bold = True
for i, row in enumerate(wf_rows, start=1):
    h8_sd_pct = float(row['lstm_h8_seed_sd']) / float(row['lstm_h8_mae']) * float(row['lstm_h8_mape'])
    h32_sd_pct = float(row['lstm_h32_seed_sd']) / float(row['lstm_h32_mae']) * float(row['lstm_h32_mape'])
    vals = [
        WF_LABELS.get(list(row.values())[0], list(row.values())[0]),
        f"{float(row['n']):.0f}",
        f"{float(row['folds']):.0f}",
        f"{float(row['ets_mape']):.2f}",
        f"{float(row['arima_mape']):.2f}",
        f"{float(row['lstm_h8_mape']):.2f} ± {h8_sd_pct:.2f}",
        f"{float(row['lstm_h32_mape']):.2f} ± {h32_sd_pct:.2f}",
    ]
    for j, v in enumerate(vals):
        t3.cell(i, j).text = str(v)
for row in t3.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for r in p.runs:
                r.font.size = Pt(9)
doc.add_paragraph().add_run("SD = standard deviation across five random LSTM initializations, rescaled from the raw mean absolute error "
             "to percentage-error units by the same factor as the reported MAPE; with only 5 seeds this SD is itself an "
             "imprecise estimate of seed-to-seed variability (relative standard error ≈ 35%) and should be read as "
             "indicative, not a precise characterization.").italic = True
doc.add_paragraph()

# Table 4 (order-selection sensitivity, all 21 series)
cap = doc.add_paragraph()
r = cap.add_run("Table 4. ")
r.bold = True
cap.add_run("Order-selection sensitivity: a uniform ARIMA(1,1,1) specification versus each series' own AICc-selected order, "
             "fit on the same scale (log or raw) used for that series' final production forecast (Methods). "
             "Tie = ΔAICc < 2 (comparable support for both specifications).")

order_cmp = {}
with open(ROOT + r"\outputs\data\forecast_order_comparison.csv", encoding="utf-8") as f:
    for row in csv.DictReader(f):
        order_cmp[row["indicator"]] = row
margin_data = {}
with open(ROOT + r"\outputs\data\aicc_margin_vs_uniform111.csv", encoding="utf-8") as f:
    for row in csv.DictReader(f):
        margin_data[row["indicator"]] = row

t4_headers = ["Indicator", "Order used", "ΔAICc vs. (1,1,1)", "Fixed-(1,1,1) 2030", "Own-order 2030", "% change", "Tie"]
t4 = doc.add_table(rows=1 + len(ROW_ORDER), cols=len(t4_headers))
t4.style = "Light Grid Accent 1"
for j, h in enumerate(t4_headers):
    r = t4.cell(0, j).paragraphs[0].add_run(h)
    r.bold = True
for i, ind in enumerate(ROW_ORDER, start=1):
    oc = order_cmp[ind]
    mg = margin_data[ind]
    is_count = ind in ("total_ncd_deaths", "hiv_new_infections_n")
    delta = float(mg["delta_aicc"])
    near_tie = mg["near_tie"].strip().lower() == "true"
    label = LABELS[ind] + (" [a]" if ind in NON_CONVERGENT else "")
    vals = [
        label,
        oc["order_used"].replace("(", "").replace(")", "").replace(" ", ""),
        f"{delta:.2f}",
        fmt_int(oc["fixed_111_2030"]) if is_count else fmt(oc["fixed_111_2030"]),
        fmt_int(oc["aic_optimal_2030"]) if is_count else fmt(oc["aic_optimal_2030"]),
        f"{float(oc['pct_change']):.1f}",
        "Yes" if near_tie else "No",
    ]
    for j, v in enumerate(vals):
        t4.cell(i, j).text = str(v)
for row in t4.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for r in p.runs:
                r.font.size = Pt(9)
doc.add_paragraph().add_run("[a] For this series, the fixed-(1,1,1) refit used as the comparison baseline itself failed to converge "
             "(Methods); its ΔAICc and % change are reported for transparency but should be read with added caution, since "
             "they compare a valid, converged own-order fit against an unreliable baseline.").italic = True
doc.add_paragraph()

# Table 5 (structural-break sensitivity)
cap = doc.add_paragraph()
r = cap.add_run("Table 5. ")
r.bold = True
cap.add_run("Structural-break sensitivity: AICc improvement from adding a level-shift dummy at the 2020 COVID-19 onset or "
             "the 2022 Ghana currency-crisis onset to each series' own AICc-selected ARIMA specification, versus the "
             "no-intervention baseline. Positive values favour the intervention model; * marks a decisive improvement "
             "(ΔAICc > 2). NT = not testable (fewer than 2 post-break or 5 pre-break observations). Post-break "
             "observation count shown in parentheses.")

sb_rows = {}
with open(ROOT + r"\outputs\data\structural_break_sensitivity.csv", encoding="utf-8") as f:
    for row in csv.DictReader(f):
        sb_rows[row["indicator"]] = row

def fmt_break_cell(testable, delta, n_post):
    if testable.strip().lower() != "true":
        return "NT"
    d = float(delta)
    marker = "*" if d > 2 else ""
    return f"{d:.1f}{marker} ({int(float(n_post))})"

t5_headers = ["Indicator", "2020 COVID-19 break: ΔAICc (n post)", "2022 currency-crisis break: ΔAICc (n post)"]
t5 = doc.add_table(rows=1 + len(ROW_ORDER), cols=len(t5_headers))
t5.style = "Light Grid Accent 1"
for j, h in enumerate(t5_headers):
    r = t5.cell(0, j).paragraphs[0].add_run(h)
    r.bold = True
for i, ind in enumerate(ROW_ORDER, start=1):
    sb = sb_rows[ind]
    vals = [
        LABELS[ind],
        fmt_break_cell(sb["covid_2020_testable"], sb["covid_2020_delta_aicc"], sb["covid_2020_n_post"]),
        fmt_break_cell(sb["currency_2022_testable"], sb["currency_2022_delta_aicc"], sb["currency_2022_n_post"]),
    ]
    for j, v in enumerate(vals):
        t5.cell(i, j).text = str(v)
for row in t5.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for r in p.runs:
                r.font.size = Pt(9)
doc.add_paragraph()

# Figure 1 (workflow diagram)
fig1_path = MS + r"\figures\figure1_workflow_diagram.png"
doc.add_picture(fig1_path, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
cap = doc.add_paragraph()
r = cap.add_run("Figure 1. ")
r.bold = True
cap.add_run("Reproducible workflow from raw public-domain exports to reliability-flagged 2030 forecasts. Step 1 (the "
             "source-provenance audit) is the reusable data-quality contribution; Steps 2-3 are the forecasting "
             "protocol applied uniformly across eligible series.")

doc.add_paragraph()
fig2_path = MS + r"\figures\figure2_aicc_margin.png"
doc.add_picture(fig2_path, width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
cap = doc.add_paragraph()
r = cap.add_run("Figure 2. ")
r.bold = True
cap.add_run("Order-selection audit: ΔAICc between each series' own selected order and a uniform ARIMA(1,1,1) refit on the "
            "same data, sorted ascending (Table 4). Bars left of the threshold are near-ties; bars beyond it indicate the "
            "series-specific order is decisively preferred. For under-five mortality, both HIV series, and PM2.5 exposure, "
            "the fixed-(1,1,1) baseline itself failed to converge (Table 4, footnote a); these bars should be read with "
            "added caution.")

doc.add_paragraph()
fig3_path = MS + r"\figures\figure3_forecast_pct_change.png"
doc.add_picture(fig3_path, width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
cap = doc.add_paragraph()
r = cap.add_run("Figure 3. ")
r.bold = True
cap.add_run("Change in the 2030 point forecast attributable to per-series order selection versus a uniform ARIMA(1,1,1), "
             "by indicator (Table 4). Most series are insensitive; a minority shift substantially.")

doc.add_paragraph()
fig4_path = MS + r"\figures\figure4_walkforward_validation.png"
doc.add_picture(fig4_path, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
cap = doc.add_paragraph()
r = cap.add_run("Figure 4. ")
r.bold = True
cap.add_run("Walk-forward validation of classical methods versus LSTM on two representative series (Table 3). MAPE = "
             "mean absolute percentage error, seed-averaged across five LSTM initializations per fold and hidden-layer size.")

doc.add_heading("Discussion", level=1)
add_md_section(doc, MS + r"\discussion.md", cite_map, skip_first_h1=True)

doc.add_heading("Limitations", level=1)
add_md_section(doc, MS + r"\limitations.md", cite_map, skip_first_h1=True)

doc.add_heading("Conclusion", level=1)
add_md_section(doc, MS + r"\conclusion.md", cite_map, skip_first_h1=True)

# Declarations
doc.add_heading("Declarations", level=1)

p = doc.add_paragraph()
p.add_run("Ethics. ").bold = True
p.add_run("Secondary analysis of public-domain, aggregate, de-identified WHO and World Bank data; no individual-level "
          "data were accessed and no ethics approval was sought or required.")

p = doc.add_paragraph()
p.add_run("Author contributions (CRediT). ").bold = True
p.add_run("Valentine Golden Ghanem: Conceptualization, Data curation, Formal analysis, Methodology, Software, "
          "Validation, Visualization, Writing – original draft, Writing – review & editing.")

p = doc.add_paragraph()
p.add_run("Competing interests. ").bold = True
p.add_run("The author declares no competing interests.")

p = doc.add_paragraph()
p.add_run("Data availability. ").bold = True
p.add_run("Processed panel files, forecast outputs, and analysis code are available at "
          "https://github.com/valentineghanem-bit/disease-burden-forecasting-ghana. A frozen, timestamped archival "
          "snapshot of the source exports (Zenodo/OSF) is in preparation and its DOI will be added here upon deposit.")

p = doc.add_paragraph()
p.add_run("Funding. ").bold = True
p.add_run("None declared.")

# References
doc.add_heading("References", level=1)
for k in cite_order:
    p = doc.add_paragraph()
    p.add_run(f"{cite_map[k]}. {BIB[k]}")

doc.core_properties.title = TITLE
doc.core_properties.author = "Valentine Golden Ghanem"
doc.core_properties.subject = "Reproducible workflow, source-provenance audit, Ghana, time-series forecasting"

out_path = MS + r"\Ghanem_2026_Ghana_Burden_Forecasting_2030.docx"
doc.save(out_path)
print("Saved:", out_path)
print("Total paragraphs:", len(doc.paragraphs))
print("Total tables:", len(doc.tables))
print("Total citations:", len(cite_order))
